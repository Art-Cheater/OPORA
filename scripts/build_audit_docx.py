"""Сборка DOCX-отчёта аудита «Опора». Запуск: python scripts/build_audit_docx.py"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "OPORA_AUDIT_FULL_REPORT.docx"


def set_run_font(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, bold=True, size={1: 18, 2: 14, 3: 12}.get(level, 11))


def p(doc, text, *, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        set_run_font(run, size=11)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, size=10)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=9)
    run.font.name = "Consolas"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(56)
    section.right_margin = Pt(56)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Полный технический аудит проекта «Опора»")
    set_run_font(run, bold=True, size=20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Только анализ · код не изменялся · состояние на HEAD 1099d33 / RELEASE 20260826h"
    )
    set_run_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))

    p(
        doc,
        "Документ предназначен для передачи другому разработчику (в т.ч. ChatGPT): "
        "полное понимание архитектуры, БД, прав, модулей, рисков и порядка работ.",
    )
    p(
        doc,
        "Важно: локально main ahead 1 относительно origin/main (9da1c36). "
        "На сервере без git push релиз 1099d33 / миграция 035 / hardening могут отсутствовать. "
        "В working tree много M-файлов (часто CRLF) — не коммитить всё подряд.",
        bold=True,
    )

    # --- 1 ---
    h(doc, "1. Общая информация о проекте", 1)
    table(
        doc,
        ["Параметр", "Факт"],
        [
            ["Назначение", "КИС: диспетчеризация заявок, объекты наружного освещения (Киров), проекты/торги/контракты, договора на опоры, почтовые обращения, импорт ЕИС, мессенджер, личные документы, роли"],
            ["Тип", "Монолитное server-rendered веб-приложение + SPA-навигация"],
            ["Python", "3.12 (Dockerfile: python:3.12-slim)"],
            ["Backend", "Flask 3.1.0, SQLAlchemy 3.1.1, Flask-Migrate 4.0.7, Flask-Login 0.6.3, Flask-WTF 1.2.2, bcrypt 4.2.1, gunicorn 23.0.0"],
            ["Frontend", "Jinja2, Bootstrap 5, vanilla JS, Leaflet (договора), OSM iframe (заявки)"],
            ["БД", "PostgreSQL 17 (Docker); опционально SQLite (USE_SQLITE=1)"],
            ["Auth", "Flask-Login + сессии + CSRF; RBAC + field-level"],
            ["Миграции", "Alembic через Flask-Migrate, head 035_integrity_indexes"],
            ["Docker", "db, web, nginx, eis-sync, inquiry-sync, documents-notify"],
            ["Внешние сервисы", "Nominatim (адреса), zakupki.gov.ru (ЕИС), IMAP mail.ru (обращения)"],
            ["Запуск", "run.py / wsgi:app; prod: gunicorn + nginx :5000→80"],
            ["Деплой", "scripts/deploy.sh; entrypoint: flask db upgrade"],
            ["Регистрация", "Публичной нет — сотрудников создаёт HR/admin в /employees"],
        ],
    )

    # --- 2 ---
    h(doc, "2. Структура проекта", 1)
    code_block(
        doc,
        """OPORA/
├── app/                    # Flask-приложение
│   ├── __init__.py         # factory, SPA, security headers, CLI, error handlers
│   ├── config.py / release.py / extensions.py
│   ├── core/               # permissions, decorators, address, upload, audit, http
│   ├── models/             # ~45 SQLAlchemy-моделей
│   ├── modules/            # blueprints
│   ├── integrations/       # zakupki (ЕИС)
│   ├── seed/               # справочники, security catalog
│   ├── static/ · templates/
├── migrations/versions/    # 001…035
├── tests/ · docker/ · scripts/ · docs/
├── docker-compose.yml · Dockerfile · requirements.txt
├── run.py · wsgi.py · .env.example""",
    )
    p(doc, "Типичный модуль: blueprint.py → routes.py → services.py → repositories.py → forms.py → templates/.")
    p(doc, "Модули: auth, main, requests, objects, projects, tenders, contracts, contractors, agreements, inquiries, eis, messenger, documents, employees, roles, field_builder, positions, reports, audit, search, notifications; _template не зарегистрирован.")

    # --- 3 ---
    h(doc, "3. Архитектура приложения", 1)
    code_block(
        doc,
        """Browser
  → nginx:80 (static / proxy)
  → gunicorn (wsgi:app)
  → Flask before_request:
       enforce_user_access (blocked/inactive)
       SPA: X-Opora-Nav=1 → g.spa_nav
  → @login_required + @permission_required / @admin_required
  → routes.py (forms / AJAX)
  → services.py (бизнес-логика, ValidationError)
  → repositories.py / SQLAlchemy → PostgreSQL
  → AuditService (по событиям)
  → Response: HTML | spa_shell partial | ajax_ok/ajax_error JSON
  → after_request: security headers, SPA headers""",
    )
    p(doc, "Цепочка закупок:")
    code_block(
        doc,
        """WorkObject → Project → TenderApplication ↔ Project (M:N)
                    → Contract ↔ WorkObject / Contractor
Request.project_id — опционально; автосвязи Request→Object нет""",
    )
    p(doc, "Workers: flask eis-sync --loop, inquiry-sync --loop, documents-notify --loop.")

    # --- 4 ---
    h(doc, "4. База данных", 1)
    bullets(
        doc,
        [
            "СУБД: PostgreSQL 17 (schema opora по умолчанию); soft-delete почти везде (deleted_at).",
            "~45 моделей в app/models/ (реестр: app/models/__init__.py).",
            "Head миграций: 035_integrity_indexes ← … ← 001_initial (линейная цепочка, 35 ревизий).",
        ],
    )
    h(doc, "4.1. Схема связей", 2)
    code_block(
        doc,
        """User
 ├── UserRole → Role → RolePermission → Permission → SystemModule
 │                    → RoleFieldPermission → FieldDefinition
 ├── Position, LoginLog, Notification, Comment, Message
 ├── MessengerConversation ↔ MessengerMessage
 └── PersonalContract → Attachment

WorkObject
 ├── Project (object_id RESTRICT)
 │    ├── Request (project_id SET NULL)
 │    ├── ProjectMember / Document / History
 │    ├── TenderProject → TenderApplication
 │    └── Contract
 ├── TenderApplication (object_id)
 └── ContractObject → Contract → ContractContractor → Contractor

Inquiry (IMAP) · PoleAgreement → PoleAgreementSite
EisImportRun → EisImportEvent
CustomField → CustomFieldValue (EAV)
AuditLog (immutable soft_delete)""",
    )
    h(doc, "4.2. Drift model ↔ migrations", 2)
    table(
        doc,
        ["Проблема", "Где"],
        [
            ["Колонка-призрак requests.assignee_id", "Создана в 001, скопирована в responsible_id в 004, не удалена; в модели нет"],
            ["Partial unique EIS/EAV/junction", "В моделях + 035; на сервере без 035 — старые absolute unique / без EIS unique"],
            ["Unique только в моделях", "contractors INN, custom_fields module+code, field_options, fields module+code, role_field_permissions, request_dispatchers name"],
            ["Soft-delete ≠ SQL CASCADE", "CASCADE только на hard delete; операционные удаления — soft"],
        ],
    )

    # --- 5 ---
    h(doc, "5. Авторизация и пользователи", 1)
    table(
        doc,
        ["Тема", "Реализация"],
        [
            ["Регистрация", "Нет публичной; EmployeeService.create_employee / seed admin"],
            ["Вход", "POST /auth/login → AuthService.authenticate"],
            ["Выход", "POST /auth/logout (+ CSRF); GET — страница подтверждения"],
            ["Пароли", "bcrypt rounds=12; legacy Werkzeug rehash"],
            ["Сессии", "Flask-Login; HttpOnly; SameSite=Lax; Secure по умолчанию False"],
            ["Роли", "admin, director, dispatcher, master, executor"],
            ["Permissions", "{module}.{action} + admin *"],
            ["Field-level", "RoleFieldPermission (NONE/VIEW/EDIT)"],
        ],
    )
    table(
        doc,
        ["Вопрос", "Где решается"],
        [
            ["Открыть страницу", "@login_required + @permission_required на route"],
            ["Изменить объект", "Тот же decorator на POST + service ValidationError"],
            ["Изменить поле", "resolve_field / can_edit_field в form→payload"],
            ["Удалить", "permission delete + service rules (admin для wipe)"],
            ["Чужой документ/чат", "get_own / ensure_access в services"],
        ],
    )
    p(doc, "После 1099d33: не-админ не может назначить admin / править админа (EmployeeService._assert_privileged_changes).", bold=True)

    # --- 6 ---
    h(doc, "6. Основные бизнес-модули", 1)
    table(
        doc,
        ["Модуль", "Назначение", "Права", "Состояние"],
        [
            ["requests", "Диспетчерские заявки, workflow, материалы, файлы, адрес/карта", "requests.*, dispatch, approve", "Работает"],
            ["objects", "Лоты освещения, Excel, автоцепочка Project/Tender/Contract", "objects.*", "Работает"],
            ["projects", "Проекты работ", "projects.*", "Работает"],
            ["tenders", "Заявки на торги", "tenders.*", "Работает"],
            ["contracts", "Контракты", "contracts.*", "Работает"],
            ["contractors", "Справочник ИНН", "contractors.*", "Работает"],
            ["agreements", "Договора на опоры + Leaflet", "agreements.*", "Работает"],
            ["inquiries", "IMAP-почта", "inquiries.*", "Работает (нужны credentials)"],
            ["eis", "zakupki import", "eis.*", "Работает (нужна сеть/TLS)"],
            ["messenger", "Чаты", "messenger.use", "Работает"],
            ["documents", "Личные файлы", "documents.use", "Работает"],
            ["employees", "Пользователи", "users.*", "Работает + hardening"],
            ["roles / field_builder / positions", "RBAC / поля", "roles.*", "Работает"],
            ["reports / audit / search", "Отчёты, журнал, поиск", "соответствующие", "Работает"],
            ["notifications", "JSON unread API", "login", "Частично (нет полноценной UI)"],
            ["materials", "—", "в каталоге прав", "Модуля нет"],
            ["_template", "заготовка", "—", "Не в registry"],
        ],
    )
    p(doc, "Отдельного модуля «аварийная бригада» нет — это статус/действие в заявках (emergency_dispatched).")

    # --- 7 ---
    h(doc, "7. Заявки — жизненный цикл", 1)
    p(doc, "Источник: app/modules/requests/workflow.py")
    code_block(
        doc,
        """new
  → emergency_dispatched | accepted_by_master | cancelled
emergency_dispatched
  → accepted_by_master | cancelled
accepted_by_master
  → completed | in_progress | cancelled
in_progress
  → completed | cancelled
completed / cancelled → конец""",
    )
    table(
        doc,
        ["Действие", "Права"],
        [
            ["Бригада выехала / передать мастеру / отмена", "requests.dispatch"],
            ["Принять / начать работу / выполнено", "requests.approve (+ edit для complete)"],
        ],
    )
    p(doc, "История: RequestHistory. Материалы: RequestMaterial. Файлы: Attachment.")
    h(doc, "7.1. Автосоздание проекта при ТЗ/ЛСР", 2)
    bullets(
        doc,
        [
            "Это НЕ заявки, а объекты.",
            "Константа ObjectService.AUTO_PROJECT_RESULT = «Обследование проведено, ТЗ подготовлено, локально-сметный расчет готов.» (objects/services.py:28–30).",
            "При срабатывании: ProjectService.create_project + WorkObject.status = in_project (objects/services.py ~251–275; projects/services.py ~292–295).",
            "Тест: tests/test_object_project_automation.py.",
            "Request сам объект/проект не создаёт.",
        ],
    )

    # --- 8 ---
    h(doc, "8. Адреса и география", 1)
    table(
        doc,
        ["Есть", "Нет / ограничено"],
        [
            ["Nominatim + heuristic каталог Кирова (app/core/address/)", "Геокодинг WorkObject (только address/name)"],
            ["На Request: address, normalized, region, district, settlement, street, house, lat/lng", "Отдельная карта объектов/проектов"],
            ["Район из OSM city_district/suburb", "Браузерный Nominatim (намеренно нет)"],
            ["Карта заявки: OSM embed (request-detail.js)", "—"],
            ["Карта договоров: Leaflet (agreements-map.js)", "—"],
            ["CLI repair-request-districts", "—"],
        ],
    )

    # --- 9 ---
    h(doc, "9. Frontend", 1)
    bullets(
        doc,
        [
            "Layouts: base.html, app.html, spa_shell.html, auth.html",
            "SPA: main.js + X-Opora-Nav / X-Opora-Partial",
            "Списки: opora-list.js + AJAX tables",
            "CSRF: meta → X-CSRFToken",
            "Модалки: crud_modals.html + module partials",
            "Тур: tour.js / tour.css",
            "Дублирование: похожие CRUD-шаблоны по модулям (ожидаемо для Jinja-монолита)",
        ],
    )

    # --- 10 ---
    h(doc, "10. Backend (качество)", 1)
    p(doc, "Крупные файлы: requests/services.py (~43KB), requests/routes.py (~39KB), objects/services.py (~38KB), eis/services.py, contracts/*, inquiries/services.py, app/__init__.py.")
    bullets(
        doc,
        [
            "Много except Exception в sync/geocode/import",
            "EIS/IMAP ручной sync через daemon thread в gunicorn worker",
            "Soft-delete + фильтры: ошибки, если забыть deleted_at",
            "N+1 на списках в целом контролируется (noload/joinedload)",
        ],
    )

    # --- 11 ---
    h(doc, "11. Безопасность", 1)
    h(doc, "11.1. Закрыто в 1099d33 (если задеплоено)", 2)
    table(
        doc,
        ["Было", "Где сейчас"],
        [
            ["Escalation → admin", "employees/services.py:_assert_privileged_changes"],
            ["resolve_field на create", "permission_service.py:resolve_field → None"],
            ["Wipe без admin", "objects/routes.py:wipe + @admin_required"],
            ["Path traversal downloads", "upload_utils.resolve_storage_path"],
            ["Logout CSRF GET", "POST + confirm"],
            ["EIS TLS off", "verify по умолчанию; EIS_SSL_VERIFY=0"],
            ["Partial unique links", "migration 035"],
        ],
    )
    h(doc, "11.2. Остаётся OPEN", 2)
    table(
        doc,
        ["Уровень", "Проблема", "Где", "Почему"],
        [
            ["HIGH", "file_upload/file_delete в каталоге, на routes не проверяются", "security_catalog vs routes (*.edit)", "Ложное ощущение контроля"],
            ["HIGH", "roles.manage может собрать почти полный набор прав", "roles/services.py", "Эскалация без кода admin"],
            ["MEDIUM", "Нет rate-limit / lockout login", "auth/services.py", "Password spray"],
            ["MEDIUM", "SESSION_COOKIE_SECURE=False default", "config.py", "Утечка cookie при HTTPS без .env"],
            ["MEDIUM", "Soft-delete проекта не всегда сбрасывает статус объекта", "projects/services.py", "Залипший in_project"],
            ["MEDIUM", "CLI wipe без admin gate", "app/__init__.py CLI", "Нужен доступ к хосту"],
            ["LOW", "Слабый CSP", "app/__init__.py", "Inline scripts остаются"],
            ["LOW", "Inquiry purge без resolve_storage_path", "inquiries/services.py ~153", "Defense-in-depth"],
            ["LOW", "Nginx без TLS", "docker/nginx.conf", "LAN-модель"],
        ],
    )
    p(doc, "Documents/messenger IDOR в текущем коде выглядят закрытыми (get_own / ensure_access).")

    # --- 12 ---
    h(doc, "12. Обработка ошибок", 1)
    bullets(
        doc,
        [
            "Глобально: 404/400/413/500/CSRF в app/__init__.py — AJAX → JSON, иначе шаблон; 500 делает db.session.rollback",
            "Доменные: ValidationError, NotFoundError, AuthenticationError",
            "Риск: необработанные исключения в тяжёлых sync; timeout Nominatim",
        ],
    )

    # --- 13 ---
    h(doc, "13. Тестирование", 1)
    p(doc, "pytest 8.3.5. CSRF в тестах выключен (conftest.py).")
    bullets(
        doc,
        [
            "Security: test_security_hardening.py",
            "Smoke/CRUD: test_http_smoke, test_core_flows, test_list_shell",
            "Requests: test_request_workflow, test_request_filters, test_wipe_requests",
            "Procurement: test_contracts, test_procurement_chain, test_object_project_automation, test_eis_*",
            "Address: test_address_*",
            "Other: documents, inquiries, agreements, roles, search, messenger perf",
        ],
    )
    p(doc, "Прогон аудита: PASS (exit 0), FAIL=0, ERROR=0. Ранее полный прогон — 156 passed.", bold=True)
    p(doc, "Дыры: rate-limit login, CSRF-on, Secure cookies, enforce file_*, E2E messenger IDOR, реальный IMAP/EIS.")

    # --- 14 ---
    h(doc, "14. Docker и запуск", 1)
    code_block(
        doc,
        """.env → docker compose up
  db (postgres:17, volume postgres_data)
  web (migrate + gunicorn) volume uploads_data
  nginx :5000→80
  eis-sync / inquiry-sync / documents-notify

Локально: USE_SQLITE=1 + flask init-db или Postgres localhost""",
    )
    p(doc, "Замечания: healthcheck web — TCP :5000, не /health; нет TLS на nginx.")

    # --- 15 ---
    h(doc, "15. Git", 1)
    code_block(
        doc,
        """branch: main
HEAD:   1099d33 [фикс] Аудит: privilege escalation…
origin: 9da1c36  → local ahead 1
dirty:  много M (подозрение на CRLF)
.gitignore: .env, instance/, *.db, docker-data/ — OK""",
    )

    # --- 16 ---
    h(doc, "16. Зависимости", 1)
    p(doc, "Только requirements.txt (нет package.json как основного). Нет Flask-Limiter. Версии pin’ами. locust / openpyxl / pymupdf / pytesseract — нагрузка, Excel, OCR.")

    # --- 17 ---
    h(doc, "17. Поиск маркеров", 1)
    bullets(
        doc,
        [
            "TODO/FIXME: не найдено",
            "Bare except:: не найдено",
            "Широкие except Exception: sync/import/geocode/upload/cache",
            "Каталог materials без модуля — мёртвое право",
            "_template не в registry",
        ],
    )

    # --- 18 ---
    h(doc, "18. Что сейчас реально работает", 1)
    h(doc, "Работает", 2)
    p(doc, "Заявки (workflow), объекты+автопроект по ТЗ/ЛСР, проекты/торги/контракты/подрядчики, договора+карта, мессенджер, документы (изоляция), роли/поля, поиск, аудит, отчёты, SPA-nav, seed admin, Docker-стек, pytest suite.")
    h(doc, "Частично", 2)
    p(doc, "Обращения/ЕИС (сеть и credentials); notifications (API без UI); геокод объектов; file_upload права в UI.")
    h(doc, "Не работает / риск без деплоя", 2)
    p(doc, "На production без push 1099d33 и migrate 035 — старые CRITICAL (escalation, wipe, path, unique) могут ещё быть живы.")
    h(doc, "Не реализовано", 2)
    p(doc, "Публичная регистрация; модуль materials; Celery; HTTPS из коробки; rate-limit login; геокарта объектов.")
    h(doc, "Неизвестно без ручной проверки", 2)
    p(doc, "Реальный IMAP, качество ЕИС на живых страницах, Nominatim с прод-IP, нагрузка gunicorn.")

    # --- 19 ---
    h(doc, "19. Главные проблемы (рейтинг)", 1)
    table(
        doc,
        ["Приоритет", "Проблема", "Где", "Почему", "Что сломается"],
        [
            ["P0", "1099d33 не на origin/сервере", "git", "Сервер без hardening", "Escalation/wipe/path на проде"],
            ["P0", "Миграция 035 может быть не применена", "Alembic", "Unique/junction", "IntegrityError re-link, дубли EIS"],
            ["P1", "Dead file_upload/file_delete", "catalog vs routes", "Ложные права", "Неверный ACL"],
            ["P1", "roles.manage почти god-role", "roles/services", "Эскалация", "Полный доступ без admin"],
            ["P1", "Нет login rate-limit", "auth/services", "Spray", "Взлом паролей"],
            ["P2", "Secure cookie default off", "config", "HTTPS misconfig", "Session theft"],
            ["P2", "assignee_id leftover", "DB", "Drift", "Мусор схема"],
            ["P2", "Project delete ↔ object status", "projects", "Soft-delete", "Залипшие статусы"],
            ["P2", "Missing unique migrations", "models vs alembic", "Дубликаты", "Целостность"],
            ["P3", "Слабый CSP / CRLF noise", "__init__ / git", "Качество", "Шум, XSS-окно"],
        ],
    )

    # --- 20 ---
    h(doc, "20. Архитектурные проблемы", 1)
    bullets(
        doc,
        [
            "Структура логична (blueprint / service / repository) — легко добавлять модули.",
            "Опасно менять: requests/*, objects/services, eis/services, permission_service, миграции junction.",
            "Связанность высокая в Object→Project→Tender→Contract.",
            "Техдолг: огромные routes/services, soft-delete vs CASCADE, materials без модуля.",
            "Спокойный рефакторинг: templates/CSS, мелкие JS, docs.",
        ],
    )

    # --- 21 ---
    h(doc, "21. Что лучше не трогать без нужды", 1)
    bullets(
        doc,
        [
            "Машину статусов заявок (workflow.py + seed статусов)",
            "Автоцепочку объектов (ObjectService._ensure_chain_for_result)",
            "PermissionService / seed security catalog без плана миграции прав",
            "Живые Alembic 001–034 на проде (только forward 035+)",
            "Volume PostgreSQL / uploads",
            "Парсер zakupki и IMAP без стендов",
            "Массовый rewrite на другой стек",
        ],
    )

    # --- 22 ---
    h(doc, "22. Рекомендуемый порядок дальнейшей работы", 1)
    code_block(
        doc,
        """Этап 1 — Выкат hardening
  Push 1099d33 → deploy.sh → Ctrl+F5 → /health=20260826h → db upgrade 035
  Риск: средний; зависимость: git remote

Этап 2 — Безопасность остатка
  Rate-limit login; Secure cookies за HTTPS; enforce или удалить file_*
  Cap roles.manage

Этап 3 — БД drift
  Drop assignee_id; миграции unique INN/custom fields; фильтры soft-delete

Этап 4 — Цепочка закупок
  Project delete → сброс WorkObject; тесты re-link после 035

Этап 5 — Заявки (полировка)
  Явная связь Request↔Object при необходимости; не ломать workflow

Этап 6 — Адреса
  Геокод объектов / единая карта — только по продуктовой нужде

Этап 7 — Frontend
  CSP, EOL, tour cache

Этап 8 — Тесты
  CSRF-on smoke; messenger IDOR; login throttle; migrate path CI

Этап 9 — Рефакторинг крупных services
  Только после зелёных regression""",
    )

    # --- 23 ---
    h(doc, "23. Паспорт проекта", 1)
    code_block(
        doc,
        """PROJECT: Опора

STACK: Python 3.12 · Flask 3.1 · SQLAlchemy · Alembic · Bootstrap 5 · vanilla JS
DATABASE: PostgreSQL 17 (Docker) · soft-delete · head 035_integrity_indexes
BACKEND: modular blueprints · services/repositories · gunicorn gthread
FRONTEND: Jinja + SPA nav (X-Opora-Nav) · Leaflet agreements · OSM requests
AUTH: Flask-Login · bcrypt · RBAC + field permissions · CSRF
DEPLOYMENT: docker-compose (db/web/nginx + 3 workers) · scripts/deploy.sh

MAIN MODULES:
  requests, objects, projects, tenders, contracts, contractors,
  agreements, inquiries, eis, messenger, documents, employees,
  roles, field_builder, reports, audit, search, notifications

CURRENT STATUS:
  Локально: 1099d33, RELEASE 20260826h, pytest PASS
  Remote/prod: вероятно без 1099d33 (ahead 1) — проверить /health и alembic

CRITICAL PROBLEMS (если прод без 1099d33):
  privilege escalation, wipe, path traversal, junction unique
  (на локальном HEAD — закрыты)

IMPORTANT FILES:
  app/__init__.py, app/core/permission_service.py, app/core/upload_utils.py
  app/modules/requests/workflow.py, objects/services.py
  app/modules/employees/services.py
  migrations/versions/035_integrity_indexes.py
  scripts/deploy.sh, docker-compose.yml

DANGEROUS AREAS:
  objects/projects/tenders/contracts chain · EIS import · IMAP · RBAC seed
  alembic on production DB · mass wipe CLI

NEXT PRIORITY:
  1) Push+deploy 1099d33 + migrate 035
  2) Login rate-limit + file_* RBAC alignment + roles.manage cap
  3) DB drift cleanup (assignee_id, missing uniques)""",
    )

    h(doc, "Примечание для следующего разработчика", 1)
    p(
        doc,
        "Перед правками сверь git log -1, /health release и flask db current на целевом окружении. "
        "Этот документ — снимок аудита без изменений кода.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
